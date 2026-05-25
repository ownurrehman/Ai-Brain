#!/usr/bin/env python3
"""Create a professional SEO proposal DOCX for mysportsinjury.co.uk"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "/Users/sheikhown/Library/CloudStorage/OneDrive-Personal/Rank Ray - Company/RR - Office Works/Proposals/SEO/SEO Proposal - MySportsInjury - Hamza.docx"

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    if color:
        run.font.color.rgb = color
    p.space_after = Pt(6)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.space_after = Pt(3)
    return p

def add_para(doc, text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if align:
        p.alignment = align
    p.space_after = Pt(6)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph("─" * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# ===================== COVER PAGE =====================
add_para(doc, "", bold=False)
add_para(doc, "", bold=False)
add_para(doc, "", bold=False)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SEO PROPOSAL")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("For mysportsinjury.co.uk")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sports Physiotherapy Clinic — Manchester, UK")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ===================== SECTION 1: ABOUT CLIENT =====================
add_heading(doc, "1. ABOUT THE CLIENT", level=1, color=RGBColor(0x1A, 0x3A, 0x5C))

add_para(doc, "Website: mysportsinjury.co.uk")
add_para(doc, "Industry: Sports Physiotherapy & Rehabilitation")
add_para(doc, "Location: Manchester, United Kingdom")
add_para(doc, "Representative: Hamza")

add_para(doc, "My Sports Injury is a leading sports physiotherapy clinic based in Manchester, UK. They specialize in treating sports injuries and offer services including physiotherapy, sports massage, rehabilitation programs, and recovery products. Their primary audience includes athletes, fitness enthusiasts, and individuals seeking professional physiotherapy care across the UK.")

add_hr(doc)

# ===================== SECTION 2: CURRENT SEO PERFORMANCE =====================
add_heading(doc, "2. CURRENT SEO PERFORMANCE", level=1, color=RGBColor(0x1A, 0x3A, 0x5C))

add_heading(doc, "2.1 Domain Overview (SEMrush UK Database)", level=2)

# Stats table
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Metric"
hdr_cells[1].text = "Value"
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ("Domain Rank", "#244,526"),
    ("Organic Keywords", "437"),
    ("Organic Traffic", "~1,339 visits/month"),
    ("Traffic Value", "£545/month"),
    ("Paid Advertising", "None currently running"),
]
for i, (metric, value) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = metric
    row[1].text = value

doc.add_paragraph()

add_heading(doc, "2.2 Current Keyword Rankings", level=2)

add_para(doc, "Based on SEMrush organic keyword analysis for the UK market:")

rankings = [
    ("'sports physio manchester'", "#1", "110", "✅ Maintained"),
    ("'medial tibial periostitis'", "#3", "480", "🎯 Close to top"),
    ("'lumbar lordosis'", "#11", "4,400", "🔼 High opportunity"),
    ("'thoracic back pain'", "#17", "2,400", "🔼 High opportunity"),
    ("'ankle syndesmosis'", "#8", "590", "🎯 Close to top"),
    ("'periostitis'", "#5", "720", "🎯 Close to top"),
    ("'lumbosacral lordosis'", "#9", "3,600", "🎯 Close to top"),
    ("'best exercises for inguinal hernia'", "#3", "90", "✅ Maintained"),
]

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Light List Accent 1'
hdr = table2.rows[0].cells
hdr[0].text = "Keyword"
hdr[1].text = "Position"
hdr[2].text = "Volume/mo"
hdr[3].text = "Status"
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for kw, pos, vol, status in rankings:
    row = table2.add_row().cells
    row[0].text = kw
    row[1].text = pos
    row[2].text = vol
    row[3].text = status

doc.add_paragraph()
add_para(doc, "Key Finding: The site already ranks #1 for a local branded term but has significant untapped potential for high-volume condition-related keywords (e.g., 'lumbar lordosis' at 4,400 monthly searches currently at position #11).")

add_hr(doc)

# ===================== SECTION 3: SEO STRATEGY =====================
add_heading(doc, "3. SEO OPPORTUNITIES & STRATEGY", level=1, color=RGBColor(0x1A, 0x3A, 0x5C))

add_para(doc, "Our 3-month SEO campaign will focus on five core pillars to improve visibility, traffic, and patient bookings for My Sports Injury.")

add_heading(doc, "3.1 Technical SEO Audit & Fixes", level=2)
add_bullet(doc, "Comprehensive site crawl to identify and fix technical errors")
add_bullet(doc, "Improve Core Web Vitals (page speed, LCP, CLS, FID)")
add_bullet(doc, "Mobile-first optimization across all key pages")
add_bullet(doc, "Implement LocalBusiness schema markup and medical/health schema")
add_bullet(doc, "Fix broken links, redirect chains, and orphan pages")
add_bullet(doc, "XML sitemap optimization and robots.txt review")

add_heading(doc, "3.2 On-Page Optimization", level=2)
add_bullet(doc, "Optimize existing 437+ keyword-targeted pages for higher rankings")
add_bullet(doc, "Rewrite title tags and meta descriptions (target <160 chars with keyword + LSI + brand)")
add_bullet(doc, "Improve H1 structure — ensure H1 differs from title tag on every page")
add_bullet(doc, "Strengthen internal linking structure between condition pages and services")
add_bullet(doc, "Add FAQ schema to high-traffic blog posts for featured snippet capture")
add_bullet(doc, "Image optimization: alt text, compression, WebP format where possible")

add_heading(doc, "3.3 Content Strategy", level=2)
add_bullet(doc, "Create condition-specific landing pages targeting high-volume terms")
add_bullet(doc, "Example targets: lumbar lordosis (4,400/mo), thoracic back pain (2,400/mo), sciatica, runner's knee")
add_bullet(doc, "Publish 8+ new SEO-optimized articles per month")
add_bullet(doc, "Content types: injury guides, rehabilitation programs, prevention tips, athlete success stories")
add_bullet(doc, "Align all content with AEO (Ask Engine Optimization) — answer-focused format for AI search")

add_heading(doc, "3.4 Local SEO (Manchester & UK)", level=2)
add_bullet(doc, "Google Business Profile full optimization — categories, services, photos, Q&A, posts")
add_bullet(doc, "Build consistent NAP citations across UK healthcare and local directories")
add_bullet(doc, "Create location-based landing pages: 'Sports Physio Manchester', 'Sports Injury Clinic [Area]'")
add_bullet(doc, "Review and response strategy for patient reviews")
add_bullet(doc, "Local backlink acquisition from Manchester sports clubs, gyms, and health organizations")

add_heading(doc, "3.5 Backlink Building & Authority", level=2)
add_bullet(doc, "Sports, health, and fitness industry outreach")
add_bullet(doc, "Guest posting on physiotherapy, sports medicine, and fitness blogs")
add_bullet(doc, "Local Manchester business directory submissions")
add_bullet(doc, "Target: 5–10 quality backlinks per month")
add_bullet(doc, "Competitor backlink gap analysis and replication")

add_hr(doc)

# ===================== SECTION 4: PRICING =====================
add_heading(doc, "4. PRICING & CONTRACT TERMS", level=1, color=RGBColor(0x1A, 0x3A, 0x5C))

add_heading(doc, "4.1 Initial Campaign", level=2)

pricing_data = [
    ("Monthly Investment", "£500 GBP"),
    ("Campaign Duration", "3 Months"),
    ("Total Investment", "£1,500 GBP"),
    ("Payment Terms", "Monthly in advance"),
    ("Start Date", "Upon agreement signature"),
]

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Medium Shading 1 Accent 1'
hdr = table3.rows[0].cells
hdr[0].text = "Item"
hdr[1].text = "Details"
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for item, detail in pricing_data:
    row = table3.add_row().cells
    row[0].text = item
    row[1].text = detail

doc.add_paragraph()

add_heading(doc, "4.2 What's Included (Monthly)", level=2)
add_bullet(doc, "Full technical SEO audit & ongoing fixes")
add_bullet(doc, "On-page optimization for all key pages")
add_bullet(doc, "8+ new SEO-optimized articles per month")
add_bullet(doc, "Local SEO & Google Business Profile management")
add_bullet(doc, "Monthly backlink acquisition (5–10 quality links)")
add_bullet(doc, "Weekly rank tracking & detailed monthly reporting")
add_bullet(doc, "Monthly strategy call with Hamza / team")

add_heading(doc, "4.3 Upgrade Path", level=2)
add_para(doc, "If Rank Ray delivers measurable, agreed-upon results within the initial 3-month period, the contract will automatically upgrade to a 1-Year Agreement. The annual retainer will be adjusted based on performance metrics, scope expansion, and competitive landscape.")

add_hr(doc)

# ===================== SECTION 5: EXPECTED RESULTS =====================
add_heading(doc, "5. EXPECTED RESULTS (3 MONTHS)", level=1, color=RGBColor(0x1A, 0x3A, 0x5C))

add_para(doc, "The following projections are conservative estimates based on the current domain strength, competition level in UK physiotherapy/healthcare SEO, and Rank Ray's historical performance with similar clients.")

results = [
    ("Organic Keywords", "437", "700+", "~60% increase"),
    ("Organic Traffic", "~1,339/mo", "3,000+/mo", "~125% increase"),
    ("Top 3 Rankings", "~5 keywords", "25+ keywords", "5x growth"),
    ("Top 10 Rankings", "~15 keywords", "50+ keywords", "3x growth"),
    ("Domain Authority", "Current", "+15-20%", "Authority growth"),
]

table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Light Grid Accent 1'
hdr = table4.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Current"
hdr[2].text = "3-Month Target"
hdr[3].text = "Growth"
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for metric, curr, target, growth in results:
    row = table4.add_row().cells
    row[0].text = metric
    row[1].text = curr
    row[2].text = target
    row[3].text = growth

doc.add_paragraph()

add_heading(doc, "5.1 Priority Keyword Targets", level=2)
add_bullet(doc, "'lumbar lordosis' (4,400/mo) — Target: Position #5 or better")
add_bullet(doc, "'thoracic back pain' (2,400/mo) — Target: Position #10 or better")
add_bullet(doc, "'sports physiotherapy manchester' — Target: Position #1")
add_bullet(doc, "'ankle syndesmosis' (590/mo) — Target: Position #3 or better")
add_bullet(doc, "'sports massage manchester' — Target: Top 5")
add_bullet(doc, "'physio for runners' — Target: Top 10")

doc.add_paragraph()
add_heading(doc, "5.2 Success Metrics", level=2)
add_bullet(doc, "Increased online appointment bookings via organic traffic")
add_bullet(doc, "Higher local search visibility in Manchester and surrounding areas")
add_bullet(doc, "Improved domain authority and backlink profile")
add_bullet(doc, "Measurable ROI: traffic value growth vs. SEO investment")

add_hr(doc)

# ===================== SECTION 6: WHY RANK RAY =====================
add_heading(doc, "6. WHY RANK RAY?", level=1, color=RGBColor(0x1A, 0x3A, 0x5C))

add_para(doc, "Rank Ray is a UK-focused digital marketing agency with deep expertise in healthcare, local, and sports industry SEO. We combine technical precision with content excellence to deliver measurable growth.")

add_bullet(doc, "Specialized experience in local SEO and healthcare/physiotherapy niches")
add_bullet(doc, "Proven track record with UK-based businesses")
add_bullet(doc, "Data-driven approach using SEMrush, Ahrefs, and proprietary tools")
add_bullet(doc, "Transparent monthly reporting and direct communication")
add_bullet(doc, "Focus on ROI: we measure success in bookings and revenue, not just rankings")

add_heading(doc, "Our Process", level=2)
add_bullet(doc, "Week 1–2: Full audit, competitor analysis, and strategy development")
add_bullet(doc, "Week 3–4: Technical fixes and quick-win on-page optimizations")
add_bullet(doc, "Month 2: Content production and local SEO implementation")
add_bullet(doc, "Month 3: Backlink acquisition, advanced optimizations, and performance review")
add_bullet(doc, "Ongoing: Monthly reporting, strategy refinement, and scaling")

add_hr(doc)

# ===================== SECTION 7: CONTACT =====================
add_heading(doc, "7. CONTACT & NEXT STEPS", level=1, color=RGBColor(0x1A, 0x3A, 0x5C))

add_para(doc, "We look forward to partnering with My Sports Injury and helping you dominate the Manchester and UK physiotherapy search landscape.")

add_para(doc, "Own-ur-Rehman Sheikh")
add_para(doc, "CEO, Rank Ray Digital Agency")
add_para(doc, "Email: contact@rankray.com")
add_para(doc, "Phone: +92-333-5261658")
add_para(doc, "Web: www.rankray.com")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Thank you for considering Rank Ray.")
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Save
doc.save(OUTPUT_PATH)
print(f"✅ SEO Proposal DOCX created: {OUTPUT_PATH}")
